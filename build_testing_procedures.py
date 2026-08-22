"""Build the maintained Word document of ChurchManager testing procedures."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('Documentation/ChurchManager.Testing.Procedures.docx')
BLUE = '2E74B5'; DARK = '1F4D78'; PALE = 'E8EEF5'; LIGHT = 'F4F6F9'; GRAY = '666666'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for tag, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = mar.find(qn('w:'+tag))
        if node is None: node = OxmlElement('w:'+tag); mar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); rep = OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); trPr.append(rep)

def keep(p, next_=False):
    pPr = p._p.get_or_add_pPr(); el = OxmlElement('w:keepNext' if next_ else 'w:keepLines'); pPr.append(el)

def add_page_num(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run('Page '); r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GRAY)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld)

def configure(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    sec = doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles
    n=styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11); n.font.color.rgb=RGBColor(0,0,0)
    n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.25
    for nm,size,color,before,after in [('Title',28,DARK,0,8),('Subtitle',12,GRAY,0,14),('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK,10,5)]:
        s=styles[nm]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=(nm.startswith('Heading') or nm=='Title')
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for nm in ['List Bullet','List Number']:
        s=styles[nm]; s.font.name='Calibri'; s.font.size=Pt(11); s.paragraph_format.left_indent=Inches(.375); s.paragraph_format.first_line_indent=Inches(-.188); s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.25
    if 'Procedure Step' not in styles:
        s=styles.add_style('Procedure Step',WD_STYLE_TYPE.PARAGRAPH); s.base_style=styles['Normal']; s.font.name='Calibri'; s.font.size=Pt(11); s.paragraph_format.left_indent=Inches(.22); s.paragraph_format.space_after=Pt(5); s.paragraph_format.keep_together=True
    for header_part in (sec.header, sec.even_page_header, sec.first_page_header):
        header=header_part.paragraphs[0]; header.text='CHURCHMANAGER  |  TESTING PROCEDURES'; header.style=styles['Header']; header.runs[0].font.name='Calibri'; header.runs[0].font.size=Pt(9); header.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
    for footer_part in (sec.footer, sec.even_page_footer, sec.first_page_footer):
        add_page_num(footer_part.paragraphs[0])

def callout(doc,label,text,fill=LIGHT):
    t=doc.add_table(rows=1,cols=1); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.columns[0].width=Inches(6.5)
    c=t.cell(0,0); c.width=Inches(6.5); shade(c,fill); set_cell_margins(c,120,160,120,160)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(label+'  '); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def table(doc,headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    set_repeat_table_header(t.rows[0])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); shade(c,PALE); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].width=Inches(widths[i]); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.add_run(str(val))
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def bullet(doc,text,style='List Bullet'):
    p=doc.add_paragraph(text,style=style); keep(p); return p

def step(doc,num,title,action,expected,evidence=''):
    p=doc.add_paragraph(style='Procedure Step'); keep(p)
    r=p.add_run(f'{num}. {title}'); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK)
    p.add_run(f'\nAction: {action}\nExpected: {expected}')
    if evidence: p.add_run(f'\nRecord: {evidence}')

doc=Document(); configure(doc)
p=doc.add_paragraph(style='Title'); p.add_run('ChurchManager Testing Procedures')
p=doc.add_paragraph('Repeatable verification for maintenance, releases, and recovery readiness',style='Subtitle')
table(doc,['Document control','Value'],[
    ('Purpose','Verify ChurchManager without risking live congregational records.'),
    ('Audience','Pastor, maintainer, or trusted tester familiar with normal ChurchManager workflows.'),
    ('Test level','ChurchManager module checks, application smoke tests, workflow acceptance tests, integrations, and release sign-off.'),
    ('Primary rule','Use a separate test database and disposable output folders. Never experiment against the only live copy.'),
], [1.65,4.85])
callout(doc,'Scope boundary','Accounting and Giving are active ChurchManager subsystems and must be tested only with fictitious data in the guarded test database. Giving screens, statements, imports, evidence files, backups, and reports are confidential even during testing. Never use production contributor data for development or acceptance testing.')
callout(doc,'Framework boundary','This procedure does not test JSForm modules, controls, SQL builders, navigation engine, or framework internals. JSForm is referenced only where ChurchManager calls it as a dependency. A separate JSForm test procedure will be prepared later.')

doc.add_heading('1. How to use this procedure',level=1)
doc.add_paragraph('Run the baseline suite before a release, after changes to forms or database logic, after moving ChurchManager to a different computer, and periodically as a recovery drill. Run the focused suite for the area changed, then repeat the smoke and data-integrity checks before approval.')
for x in ['Tester writes the actual result, not merely “looks good.”','A failed expected result is a defect even when a workaround exists.','Never send test email to the congregation or write test events to the production calendar.','Preserve the sermon archive: stable IDs, metadata, file links, search results, and exportability are high-priority acceptance criteria.']: bullet(doc,x)

doc.add_heading('2. Test record and result codes',level=1)
table(doc,['Field','What to record'],[
('Run ID','Date plus a short label, for example 2026-08-09-release-candidate.'),('Build tested','Commit, folder copy, or application version and database schema snapshot.'),('Environment','Computer, Windows version, Python environment, database server, and test database name.'),('Tester','Name or initials.'),('Evidence','Screenshot, log excerpt, generated document, query result, or test data ID.'),('Result','PASS, FAIL, BLOCKED, or NOT RUN, with a brief reason.')],[1.45,5.05])
table(doc,['Code','Meaning','Release treatment'],[
('PASS','Observed result matches expected result.','May proceed.'),('FAIL','Observed result differs or data is damaged.','Correct and retest.'),('BLOCKED','Dependency or environment prevented a valid test.','Resolve or explicitly accept the risk.'),('NOT RUN','Test was intentionally omitted.','Document why and who approved omission.')],[.85,2.65,3.0])

doc.add_heading('3. Safety and test-environment preparation',level=1)
step(doc,1,'Protect production','Close ChurchManager. Make a database backup and preserve the current application folder or a version-control reference.','The original data and program version can be restored.','Backup filename, size, date, and storage location.')
step(doc,2,'Create isolated databases','Restore the application data into ChurchDBTest and the framework configuration into JSFormTest. Use permissions intended for testing.','Launching test mode cannot write to the production databases.','Test database names and server.')
step(doc,3,'Verify the Windows credential','Open Windows Credential Manager > Windows Credentials and confirm a Generic Credential named ChurchManager/Test exists for the configured username. Never record its password.','The test credential exists and its username matches churchmanager.json.','Credential target and username only; never record the password.')
step(doc,4,'Use disposable outputs','Point report, Word, PDF, temporary, and backup output locations to a test folder.','Generated files are separated from the operational archive.','Test output folder.')
step(doc,5,'Disable outward communication','Use test-only email recipients. Disconnect or substitute calendar access if a test could write externally.','No congregant receives a test message and no production calendar is altered.','Test recipient and calendar policy.')
step(doc,6,'Prepare representative data','Include a family and person, future worship service, propers/readings, hymns, sermon with linked file, participant roles, prayer, announcement, and attendance event.','Every in-scope workflow has known data to exercise.','IDs or labels for the test records.')
callout(doc,'Stop condition','If any screen shows production records, any path points at the sole sermon/document archive, or recipients include real distribution lists, stop and correct the environment before continuing.','FFF2CC')
callout(doc,'Credential failure','If ChurchManager/Test is missing or its username does not match, stop and correct Windows Credential Manager. Never add a password to churchmanager.json, a batch file, an environment variable, a command line, or a test log.','FFF2CC')

doc.add_heading('4. ChurchManager static and automated preflight checks',level=1)
doc.add_paragraph('These checks cover ChurchManager-owned Python modules, configuration, application form definitions, and report assets. They do not validate or exercise JSForm source code. Run them from the ChurchManager project folder with the project’s intended Python environment.')
doc.add_paragraph('To include the optional read-only database checks, set CHURCHMANAGER_RUN_DB_TESTS=1 and run python run_churchmanager_tests.py. The runner reads the ChurchDBTest settings from churchmanager.json and retrieves the password from ChurchManager/Test in Windows Credential Manager. It must never request or accept a clear-text password.')
table(doc,['Check','Procedure','Pass criteria'],[
('Python compile','Compile changed ChurchManager Python modules without running the application. Exclude the neighboring JSForm project.','No syntax errors in ChurchManager modules.'),
('Application JSON parse','Parse every ChurchManager Forms/*.json file as application data.','All ChurchManager form files are valid JSON.'),
('Application form inventory','Check that each ChurchManager form referenced by cm.py or another ChurchManager module has a corresponding local definition or a documented dependency exception.','No missing ChurchManager form asset; no attempt is made to test the JSForm loader.'),
('Required files','Confirm cm.py, churchmanager.json, Forms, SQL, schema, report templates, and required scripts exist.','No required component is missing.'),
('Report inventory','Compare enabled test-database catalog entries with JSForm visual report definitions.','Every enabled report maps to an approved JSON starter and dataset contract.'),
('Secrets check','Review changed files, launch settings, logs, and instructions for passwords, tokens, or private exports.','No password appears in source, configuration, environment variables, command lines, logs, or packaged output.'),
],[1.2,3.3,2.0])
callout(doc,'Retired report codes','Migration 065 disables historical report codes that have no approved JSForm visual definition. Unknown codes fail closed rather than invoking an external report engine.')

doc.add_heading('5. Startup and smoke test',level=1)
for i,(a,e) in enumerate([
('Start ChurchManager using the test database arguments/configuration.','Main window opens without an unhandled exception.'),
('Observe startup with normal network access, then with network unavailable if practical.','Network status is handled clearly; local work is not silently damaged.'),
('Open each main-menu area that is in operational use.','Each intended form opens, lays out legibly, and closes.'),
('Navigate First, Previous, Next, and Last on a populated form.','Correct records appear without skipping, duplication, or errors.'),
('Create a disposable record, update it, close/reopen the form, then delete or retire it through the normal workflow.','Values persist correctly and related data remains consistent.'),
('Close ChurchManager normally.','Application exits without leaving an error dialog or an obviously orphaned process.')],1): step(doc,i,'Smoke check',a,e)

doc.add_heading('6. Core data-entry acceptance tests',level=1)
table(doc,['Area','Test procedure','Expected result'],[
('Church setup','Open the church record; review configured identity and paths without changing live-like values.','Record loads; path-dependent actions use test locations.'),
('Families and people','Create a test family and person; add address, contact, membership, and date details; reopen and edit one value.','Parent-child links, required fields, choices, save, navigation, and reopen all work.'),
('Validation','Leave a required field empty and try to save; enter representative dates and selections.','Invalid input is rejected clearly; valid data saves in the correct format.'),
('Linked forms','Open address/contact/date forms from a saved parent and add a row.','New child row receives the correct parent ID and appears after refresh.'),
('Files/documents','Attach or open a disposable document through the configured test path.','Correct file opens; cancel and missing-file cases do not corrupt the record.'),
('Projects/tasks/checklists','Create a disposable project/task or checklist item, change status, and reopen.','State and relationships persist as expected.'),
],[1.2,3.4,1.9])

doc.add_heading('7. Worship workflow test',level=1)
doc.add_paragraph('Use one future test service and carry the same record through the whole sequence so cross-form relationships are exercised.')
steps=[
('Create the service','Create or copy a future service with date, time, type, and location.','The service has a stable ID and reopens with the same values.'),
('Assign propers and readings','Select the appropriate propers and readings, including an alternate reading if supported.','Selections persist and appear in dependent views.'),
('Add hymns','Add opening, sermon, Communion, and closing hymn usage as applicable.','Hymn names/numbers and sequence are retained; usage history reflects the test service.'),
('Add sermon','Create a test sermon record and link a disposable outline/document.','Sermon metadata is searchable and the linked file opens.'),
('Build order of service','Generate or open the order-of-service record and components.','Expected components appear in the intended order with no missing references.'),
('Add prayers and announcements','Create date-bounded test items and select them for the service week.','Date/week rules include eligible items and exclude ineligible items.'),
('Schedule participants','Generate assignments for configured roles.','Eligible participants are assigned once; existing assignments are not duplicated.'),
('Preview notification','Generate the participant notification using test recipients only.','Recipients, service details, and assignments are accurate; no production send occurs.'),
('Record attendance','Create the attendance event and record test attendance/Communion.','Counts and individual records persist and remain tied to the correct service.'),
]
for i,x in enumerate(steps,1): step(doc,i,x[0],x[1],x[2])

doc.add_heading('8. Sermon archive regression test',level=1)
callout(doc,'Priority','The sermon archive is especially valuable. A release should not be accepted if it weakens identification, metadata, links, searchability, or recoverability.')
for x in [
'Choose several known sermons from different years and record their IDs before testing.',
'Search by date, text/theme field, Scripture reference, preacher, and service linkage where supported.',
'Open each sermon and its outline/document link; confirm the file belongs to that record.',
'Edit a harmless field in a disposable copy, reopen it, and confirm the ID did not change.',
'Exercise any rename or file-link maintenance function only against disposable files.',
'Compare sermon count and selected IDs before and after the run; investigate any unexplained change.',
'Verify that a backup or portable export contains sermon IDs, metadata, links, and enough information to rebuild search outside ChurchDB.'
]: bullet(doc,x)

doc.add_heading('9. Reports and generated documents',level=1)
doc.add_paragraph('Test report data and visual output. A successful process exit alone is not sufficient.')
table(doc,['Test','Procedure','Pass criteria'],[
('Catalog launch','Open Reports, select an enabled report, supply required parameters, and run it.','Correct template is selected and a new output file is created.'),
('Parameter handling','Try valid boundary dates/IDs and omit a required parameter.','Valid input works; missing/invalid input is handled clearly.'),
('PDF inspection','Open the generated PDF and inspect every page.','No blank unexpected pages, clipping, missing groups, broken totals, or stale data.'),
('Locked output','Keep an older PDF open and rerun if safe.','Failure is understandable and does not overwrite unrelated files.'),
('Order of service','Generate the document for the test service.','Readings, hymns, headings, and sequence match the source records.'),
('Prayers/announcements','Generate for controlled dates including week-of-month boundaries.','Only eligible items appear; today-override behavior is documented and controlled.'),
('Member directory','Generate using test data.','Names/contact fields are correct and layout remains readable.'),
],[1.2,3.3,2.0])

doc.add_heading('10. Integrations and failure handling',level=1)
table(doc,['Integration','Normal test','Failure test'],[
('Database','Read, insert, update, navigate, and reopen representative records.','Disconnect or use an invalid test endpoint; verify the error is visible and no partial data is claimed saved.'),
('Filesystem','Open and generate disposable documents.','Use a missing/read-only test path; verify a useful failure and no silent record damage.'),
('Visual reports','Generate a known report and inspect the PDF.','Use a missing or invalid test definition; verify diagnosis identifies the definition or dataset problem.'),
('Email','Preview/send only to a controlled test address.','Remove the test address or deny connection; verify no false success message.'),
('Google Calendar','Read a controlled service-week calendar if authorization is available.','Test unavailable network/expired authorization; verify the application remains understandable and recoverable.'),
('Backup','Generate a test backup, verify nonzero size, and perform a restore drill.','Use an invalid test output path; verify “complete” is not accepted without a real file.'),
],[1.2,2.65,2.65])

doc.add_heading('11. Backup and restore drill',level=1)
step(doc,1,'Create the backup','Run the configured backup against the test database.','A new, nonempty file exists at the intended test location.')
step(doc,2,'Inspect the artifact','Record filename, size, timestamp, command result/log, and optionally a checksum.','Evidence identifies exactly what will be restored.')
step(doc,3,'Restore separately','Restore into a new database name, never over the source of the drill.','Restore completes without changing the source database.')
step(doc,4,'Verify restored workflows','Open representative person, service, sermon, prayer, announcement, and attendance records; run one report.','Counts and sampled relationships match the source; linked external files are accounted for separately.')
step(doc,5,'Record recovery limits','List items not contained in the database backup, including documents, report templates, configuration, credentials, and external integrations.','Recovery instructions cover both database and required filesystem/configuration assets.')

doc.add_heading('12. Change-focused regression matrix',level=1)
table(doc,['Changed component','Minimum focused tests'],[
('Forms/*.json','JSON parse, ChurchManager reference check, and the affected ChurchManager workflow. Treat loading/navigation mechanics as JSForm out of scope unless ChurchManager customization caused the failure.'),
('cm.py','Startup/smoke, affected menu action, normal and failure path, clean close.'),
('Database query/schema','Restore test baseline, CRUD for affected forms, relationship checks, reports that consume changed fields.'),
('fnSchedule.py','Controlled service, role eligibility, duplicate prevention, assignments, test-only notification.'),
('rpt*.py','Controlled inputs, generated file existence, content comparison, every-page visual inspection.'),
('Visual report definition/catalog','Catalog mapping, approved dataset, parameters, PDF generation, grouping/totals, pagination, locked-output handling.'),
('Sermon/file handling','Stable IDs, search, link open, rename on disposable files, count and missing-file scan.'),
('Configuration/paths','Startup, all changed paths, missing-path behavior, no production crossover.'),
],[1.65,4.85])

doc.add_heading('13. Defect handling and retest',level=1)
for x in ['Give the defect a unique ID and short action-oriented title.','Record the exact test environment, starting data, steps, expected result, actual result, and evidence.','Classify impact: data loss/privacy, worship-critical, archive/recovery, major workflow, minor usability, or cosmetic.','Restore the test baseline before retesting if the failure may have changed data.','Retest the failed case, its nearest related workflows, the smoke suite, and data-integrity checks.','Do not close a defect merely because the original error disappeared; confirm the expected result and absence of side effects.']: bullet(doc,x)

doc.add_heading('14. Release acceptance checklist',level=1)
checks=['Production backup verified and rollback reference recorded.','Test configuration proven isolated from production.','Static/form checks passed or exceptions approved.','Startup and smoke suite passed.','Changed workflows and their regression matrix passed.','Worship workflow passed when affected.','Sermon archive IDs, search, and links passed.','Reports were visually inspected, not merely generated.','Backup restore drill is current.','No test email/calendar writes reached production.','No new secrets or private exports were packaged.','All FAIL, BLOCKED, and NOT RUN results have a documented disposition.']
for x in checks: bullet(doc,'☐ '+x)
table(doc,['Approval','Name / date / notes'],[('Tester',''),('Maintainer or owner',''),('Release/rollback reference','')],[2.0,4.5])

doc.add_heading('Appendix A. Test-run worksheet',level=1)
table(doc,['Test ID','Area / scenario','Result','Evidence / notes'],[(f'T-{i:02d}','','','') for i in range(1,13)],[.7,2.2,.8,2.8])

p = doc.add_heading('Appendix B. Recommended future ChurchManager automation',level=1)
p.paragraph_format.page_break_before = True
doc.add_paragraph('ChurchManager currently relies heavily on manual and exploratory testing. The following automation would provide the best protection without replacing end-to-end human verification:')
for x in ['Non-interactive parsing and reference checks for ChurchManager application form files, without testing the JSForm parser itself.','ChurchManager module tests for menu dispatch, scheduling rules, date/week selection, output-path selection, and safe error handling using substituted dependencies.','A report-catalog validator for unique codes, template existence, parameter compatibility, and output creation.','Database integration tests for ChurchManager workflows using a disposable restored fixture.','Sermon archive checks for duplicate IDs, missing files, broken links, and portable-export completeness.','Golden-output comparisons for selected ChurchManager-generated documents, with a deliberate review process for approved layout changes.']: bullet(doc,x)

doc.add_heading('Appendix C. Related project references',level=1)
for x in ['Documentation/ChurchManager.Application.md — application operation, architecture, reports, backups, troubleshooting, and maintenance.','Documentation/form.documentation.md — reference for ChurchManager-owned JSON application definitions; framework testing is excluded.','tests/ — maintained ChurchManager application tests.','Forms/ and visual_reports/definitions/ — ChurchManager screen and report definitions under test.','Separate future document — JSForm framework and module testing procedures.']: bullet(doc,x)

# Keep headings with following content and set core metadata.
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'): keep(p,True)
doc.core_properties.title='ChurchManager Testing Procedures'; doc.core_properties.subject='Repeatable ChurchManager verification and release acceptance'; doc.core_properties.author='ChurchManager Project'
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT)
print(OUT.resolve())
