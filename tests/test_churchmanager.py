"""Safe automated checks for ChurchManager-owned code and assets.

These tests deliberately do not import ChurchManager application modules because
several of them connect to the database, create a wx application, write output,
or launch another program at import time.  They also do not inspect or test the
neighboring JSForm package.
"""

from __future__ import annotations

import ast
import csv
import contextlib
import io
import json
import py_compile
import tempfile
import unittest
from datetime import date, datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FORMS = ROOT / "Forms"
VISUAL_REPORTS = ROOT / "visual_reports" / "definitions"


class TestCopyrightSensitiveWorshipFields(unittest.TestCase):
    def test_introit_and_service_psalm_fields_are_permanently_removed(self):
        migration = (ROOT / "migrations" / "037_remove_copyrighted_introit_fields.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn("DROP COLUMN IF EXISTS PsalmorIntroit", migration)
        self.assertIn("DROP COLUMN IF EXISTS Introit", migration)
        current_sources = (
            ROOT / "unified_worship_service_dialog.py",
            ROOT / "worship_service_dialog.py",
            FORMS / "frmPropers.json",
            ROOT / "visual_reports" / "report_inventory.py",
            ROOT / "visual_reports" / "definitions" / "CMWP01.json",
        )
        combined = "\n".join(path.read_text(encoding="utf-8-sig") for path in current_sources)
        self.assertNotIn("PsalmorIntroit", combined)
        self.assertNotIn('p.Introit', combined)


class TestPrintedLsbTuneMetadata(unittest.TestCase):
    def test_tune_source_is_exactly_the_printed_lsb_range(self):
        rows = list(csv.DictReader(
            (ROOT / "data" / "lsb_printed_hymn_tunes.csv").open(encoding="utf-8")
        ))
        numbers = [int(row["HymnNumber"]) for row in rows]
        self.assertEqual(len(rows), 636)
        self.assertEqual(numbers, list(range(331, 967)))
        self.assertTrue(all(row["Tune"].strip() for row in rows))
        migration = (ROOT / "migrations" / "038_add_printed_lsb_hymn_tunes.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn("ADD COLUMN IF NOT EXISTS Tune", migration)
        self.assertIn("t.HymnNumber BETWEEN 331 AND 966", migration)


class TestLsbStarterReferenceConvention(unittest.TestCase):
    def test_starter_liturgical_pages_and_psalms_are_normalized(self):
        migration = (
            ROOT / "migrations" / "064_normalize_lsb_starter_references.sql"
        ).read_text(encoding="utf-8")
        self.assertIn("template.IsStarter=1", migration)
        self.assertIn("CONCAT('LSB p. ',SUBSTRING(line.ReferenceText,5))", migration)
        self.assertIn("|330)", migration)
        self.assertIn("line.Label='Psalm'", migration)
        self.assertIn("SET line.ReferenceText=NULL", migration)
        self.assertIn("3[0-2][0-9]|330", migration)

# ChurchManager-owned operational modules. Archived code, saved copies,
# conversion utilities, and JSForm are intentionally absent.
OPERATIONAL_MODULES = (
    "cm.py",
    "fnCMargParse.py",
    "churchmanager_mode.py",
    "application_context.py",
    "authentication.py",
    "authorization.py",
    "bootstrap_test_master.py",
    "form_factory.py",
    "main_menu.py",
    "startup.py",
    "login_dialog.py",
    "user_admin.py",
    "process_service.py",
    "backup_service.py",
    "bulletin_orders.py",
    "bulletin_order_dialog.py",
    "bulletin_order_generator_dialog.py",
    "weekly_bulletin_order_dialog.py",
    "unified_worship_service_dialog.py",
    "attendance_dialog.py",
    "single_instance.py",
    "report_service.py",
    "report_support.py",
    "fnSchedule.py",
    "fnUtil.py",
    "network.py",
    "liccalendar.py",
    "rptAnnouncement.py",
    "rptPrayers.py",
    "rptMemberDirectory.py",
    "sermon2blogger.py",
)


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def load_function_without_importing(module_path: Path, function_name: str, globals_: dict):
    """Compile one pure function from a module without executing module setup."""
    tree = ast.parse(module_path.read_text(encoding="utf-8-sig"), filename=str(module_path))
    function = next(
        (node for node in ast.walk(tree) if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == function_name),
        None,
    )
    if function is None:
        raise AssertionError(f"{function_name} not found in {module_path.name}")
    isolated = ast.Module(body=[function], type_ignores=[])
    ast.fix_missing_locations(isolated)
    namespace = dict(globals_)
    exec(compile(isolated, str(module_path), "exec"), namespace)
    return namespace[function_name]


class TestWorshipPlanningStructure(unittest.TestCase):
    def test_unified_service_can_delete_only_its_weekly_order_line(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn('(\"Delete Line\", self.on_delete_line)', source)
        self.assertIn("del self.working_lines[index]", source)
        self.assertIn("The reusable template will not be changed.", source)
        self.assertIn("normalize_line_sequences(self.working_lines)", source)

    def test_weekly_hymns_store_title_as_value_and_number_as_reference(self):
        editor = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        repository = (ROOT / "bulletin_orders.py").read_text(encoding="utf-8")
        migration = (
            ROOT / "migrations" / "055_separate_weekly_hymn_title_and_number.sql"
        ).read_text(encoding="utf-8")
        self.assertIn('line["value"] = selected[2] or ""', editor)
        self.assertIn('line["reference"] = selected[1] or ""', editor)
        self.assertIn("SET WeeklyValue=?,ReferenceText=?", repository)
        self.assertIn("line.WeeklyValue", migration)
        self.assertIn("line.ReferenceText", migration)

    def test_worship_position_editor_has_a_usable_dialog_size(self):
        source = (ROOT / "worship_scheduling.py").read_text(encoding="utf-8")
        self.assertIn('title="Edit Worship Position" if row else "Add Worship Position"', source)
        self.assertIn("size=(560,330)", source)
        self.assertIn("desc.SetMinSize((-1,100))", source)
        self.assertIn('(\"Remove Assignment\",self.on_remove)', source)

    def test_deleting_custom_template_preserves_its_weekly_orders(self):
        migration = (
            ROOT / "migrations" / "054_preserve_weekly_order_snapshots.sql"
        ).read_text(encoding="utf-8")
        self.assertIn("MODIFY COLUMN TemplateID int NULL", migration)
        self.assertIn("TemplateName", migration)
        self.assertIn("ON DELETE SET NULL", migration)
        self.assertIn("ConditionType", migration)
        repository = (ROOT / "bulletin_orders.py").read_text(encoding="utf-8")
        delete_method = repository.split("def delete_custom_template", 1)[1].split("def save_line", 1)[0]
        self.assertNotIn("DELETE FROM tblServiceBulletinOrder WHERE TemplateID", delete_method)
        self.assertNotIn("DELETE l FROM tblServiceBulletinOrderLine", delete_method)

    def test_replacement_template_reuses_exact_service_hymn_selections(self):
        repository = (ROOT / "bulletin_orders.py").read_text(encoding="utf-8")
        self.assertIn("selected_hymns = cursor.fetchall()", repository)
        self.assertIn('if hymn[1] == line[5]', repository)
        self.assertIn("selected_hymns.pop(match_index)", repository)
        self.assertIn("DELETE FROM tblHymnUsage WHERE ServiceID=?", repository)

    def test_suggested_hymn_roles_use_full_liturgical_names(self):
        migration = (ROOT / "migrations" / "033_rename_suggested_hymn_roles.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn("'Hymn of Invocation'", migration)
        self.assertIn("'Hymn of the Day'", migration)
        form = json.loads(
            (ROOT / "Forms" / "frmProperHymnSuggestion.json").read_text(encoding="utf-8")
        )["frmProperHymnSuggestionFORM"]
        choices = form["CONTROLS"]["SuggestedAs"]["choices"]
        self.assertIn("Hymn of Invocation", choices)
        self.assertIn("Hymn of the Day", choices)
        self.assertIn("Distribution Hymn", choices)
        self.assertNotIn("Entrance", choices)
        self.assertNotIn("Of the Day", choices)
        self.assertNotIn("Communion", choices)

    def test_ds1_starter_has_three_distribution_hymn_slots(self):
        migration = (ROOT / "migrations" / "034_add_ds1_distribution_hymns.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn("l.Label = 'Distribution Hymn'", migration)
        self.assertIn("241,'HYMN','Distribution Hymn','SERVICE_HYMN','Distribution Hymn'", migration)
        self.assertIn("242,'HYMN','Distribution Hymn','SERVICE_HYMN','Distribution Hymn'", migration)
        self.assertNotIn("Distribution Hymn 1", migration)
        self.assertNotIn("Distribution Hymn 2", migration)
        self.assertNotIn("Distribution Hymn 3", migration)
        self.assertIn("'Distribution Hymn'", migration)

    def test_lsb_reading_roles_are_normalized_by_migration(self):
        migration = (ROOT / "migrations" / "032_normalize_lsb_reading_roles.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn("LIKE 'LSB %'", migration)
        self.assertIn("'Old Testament'", migration)
        self.assertIn("'Epistle'", migration)
        self.assertIn("'Gospel'", migration)

    def test_proper_hymn_suggestions_are_normalized(self):
        migration = (ROOT / "migrations" / "026_add_proper_hymn_suggestions.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn("CREATE TABLE IF NOT EXISTS tblProperHymnSuggestion", migration)
        self.assertIn("FOREIGN KEY (PropersID)", migration)
        self.assertIn("FOREIGN KEY (HymnID)", migration)
        self.assertIn("UNIQUE KEY uq_proper_hymn_suggestion", migration)

    def test_proper_hymn_suggestion_priority_is_removed(self):
        migration = (ROOT / "migrations" / "031_remove_proper_hymn_suggestion_priority.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn("ALTER TABLE tblProperHymnSuggestion", migration)
        self.assertIn("DROP COLUMN Priority", migration)

        form = json.loads((ROOT / "Forms" / "frmPropers.json").read_text(encoding="utf-8"))
        controls = form["frmPropersFORM"]["CONTROLS"]
        suggestion_grid = controls["dvlHymnSuggestions"]
        self.assertNotIn("Priority", suggestion_grid["table"]["fields"])
        self.assertEqual(suggestion_grid["column"][0]["lookup"]["display"], "Hymn")

    def test_church_selects_primary_hymnal(self):
        migration = (ROOT / "migrations" / "027_add_church_primary_hymnal.sql").read_text(
            encoding="utf-8"
        )
        church = load_json(FORMS / "frmChurch.json")["frmChurchFORM"]
        suggestions = load_json(FORMS / "frmProperHymnSuggestion.json")[
            "frmProperHymnSuggestionFORM"
        ]
        weekly_plan = (ROOT / "weekly_worship_plan_dialog.py").read_text(encoding="utf-8")
        self.assertIn("PrimaryHymnalID", migration)
        self.assertEqual(
            church["CONTROLS"]["PrimaryHymnalID"]["lookupchoices"]["name"], "tblHymnal"
        )
        self.assertEqual(
            suggestions["FORM"]["linkedform"]["frmHymn"]["table"]["name"], "tblHymn"
        )
        self.assertEqual(
            suggestions["CONTROLS"]["btnHymnCatalog"]["action"],
            ["openlinkedform", "frmHymn"],
        )
        self.assertIn("WHERE HymnalID=?", weekly_plan)

    def test_church_selects_an_optional_default_lectionary_edition(self):
        migration = (ROOT / "migrations" / "080_complete_lectionary_edition_cutover.sql").read_text(
            encoding="utf-8"
        )
        church = load_json(FORMS / "frmChurch.json")["frmChurchFORM"]
        control = church["CONTROLS"]["PrimaryLectionaryEditionID"]
        self.assertIn("DROP COLUMN IF EXISTS PrimaryLectionarySystemID", migration)
        self.assertIn("vwLectionaryEditionLookup", migration)
        self.assertNotIn("PrimaryLectionarySystemID", church["CONTROLS"])
        self.assertEqual(control["lookupchoices"]["name"], "vwLectionaryEditionLookup")
        self.assertTrue(control["lookupchoices"]["allowblank"])
        self.assertEqual(control["lookupchoices"]["blanklabel"], "No default edition")
        self.assertEqual(
            control["posch"][0], church["CONTROLS"]["PrimaryHymnalID"]["posch"][0]
        )
        self.assertEqual(
            control["sizech"], church["CONTROLS"]["PrimaryHymnalID"]["sizech"]
        )


class TestChurchManagerPython(unittest.TestCase):
    def test_worship_service_save_does_not_reference_removed_html_output(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertNotIn("GeneratedHtml", source)
        self.assertIn("GeneratedPlainText=NULL", source)

    def test_worship_line_actions_use_line_type_for_packaged_templates(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn('upper() == "HYMN"', source)
        self.assertIn('line_type == "READING"', source)
        self.assertIn('line_type == "HYMN"', source)
        self.assertNotIn('line["source"] != "SERVICE_HYMN"', source)

    def test_unified_worship_repository_uses_portable_connection(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn("self.connection = portable_connection(connection)", source)

    def test_unified_worship_editor_supports_line_editing_and_movement(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn("self.on_edit_line", source)
        self.assertIn("self.on_move_line(-1)", source)
        self.assertIn("self.on_move_line(1)", source)
        self.assertIn("normalize_line_sequences(self.working_lines)", source)

    def test_unified_hymn_picker_uses_safe_search_and_service_context(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn("class HymnPickerDialog", source)
        self.assertIn('columns = {', source)
        self.assertIn('LIKE ?', source)
        self.assertIn('Already used', source)
        self.assertIn('Clear This Position', source)
        self.assertIn('wx.EVT_LIST_ITEM_ACTIVATED', source)

    def test_unified_hymn_picker_sorts_on_header_double_click(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn("wx.EVT_LIST_COL_CLICK", source)
        self.assertIn("def on_column_click", source)
        self.assertIn("now - previous_time > 0.65", source)
        self.assertIn("self.rows.sort", source)

    def test_selected_proper_initializes_editable_printed_liturgical_title(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn('self.fields["liturgical"].SetValue(str(detail[3] or ""))', source)
        self.assertIn('self.fields["liturgical"].GetValue().strip() or None', source)

    def test_new_and_existing_services_both_use_unified_editor(self):
        source = (ROOT / "worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertNotIn('form_factory.create(\n            "frmService"', source)
        self.assertIn("repository.create_service(church_id)", source)
        self.assertIn("show_unified_worship_service(", source)
        self.assertIn("repository.discard_unsaved_service(new_id)", source)

    def test_service_deletion_protects_history_and_is_audited(self):
        picker = (ROOT / "worship_service_dialog.py").read_text(encoding="utf-8")
        repository = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn('(\"Delete Service\", self.on_delete)', picker)
        self.assertIn("def delete_service", repository)
        self.assertIn('dependencies.append("recorded attendance")', repository)
        self.assertIn('("tblServiceRole", "participant assignment(s)")', repository)
        self.assertIn("WORSHIP_SERVICE_DELETED", repository)
        self.assertIn("DELETE FROM tblHymnUsage WHERE ServiceID=?", repository)
        self.assertIn("DELETE FROM tblAttendanceEvent WHERE ServiceID=?", repository)

    def test_communion_checkbox_reactivates_required_template_lines(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn("control.Bind(wx.EVT_CHECKBOX, self.on_communion)", source)
        self.assertIn("def _refresh_conditional_lines", source)
        self.assertIn("BulletinOrderGenerator.condition_included", source)
        self.assertIn('line["included"] =', source)

    def test_unified_worship_save_persists_service_order_and_hymns_together(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn("def save(self, service_id, service_values, template_id, lines)", source)
        self.assertIn("UPDATE tblService SET DateTime=?", source)
        self.assertIn("DELETE FROM tblServiceBulletinOrderLine WHERE ServiceID=?", source)
        self.assertIn("INSERT INTO tblServiceBulletinOrderLine", source)
        self.assertIn("INSERT INTO tblHymnUsage", source)
        self.assertIn("self.connection.commit()", source)
        self.assertIn("self.connection.rollback()", source)
        self.assertIn('key in ("church", "os_note")', source)

    def test_service_color_override_uses_choices_and_updates_report_color(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        migration = (ROOT / "migrations" / "056_add_service_liturgical_color_override.sql").read_text(
            encoding="utf-8"
        )
        self.assertIn('choice_values("Color")', source)
        self.assertIn("Use Proper color", source)
        self.assertIn("selected = event.GetString()", source)
        self.assertIn("self.liturgical_color_swatch.Update()", source)
        self.assertIn("LiturgicalColorOverride=?", source)
        self.assertIn("LiturgicalColorOverride VARCHAR(32)", migration)
        self.assertIn("NULLIF(TRIM(s.LiturgicalColorOverride),'')", migration)

    def test_unified_worship_location_uses_tblchoices(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn('SELECT Choices FROM tblChoices WHERE Field=?', source)
        self.assertIn('choice_values("Location")', source)
        self.assertIn('key in ("proper", "sermon", "location", "color_override")', source)

    def test_unified_worship_uses_separate_native_date_and_time_pickers(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn('wx.adv.DatePickerCtrl', source)
        self.assertIn('wx.adv.TimePickerCtrl', source)
        self.assertIn('("date_time", "Date and time"', source)
        self.assertIn('label="Service date:"', source)
        self.assertIn('label="Time:"', source)
        self.assertIn('selected_date.GetYear()', source)
        self.assertIn('selected_time.GetHour()', source)

    def test_unified_worship_reserves_space_for_right_scrollbar(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8")
        self.assertIn("right_layout.AddSpacer(32)", source)

    def test_cm_has_guarded_main_entrypoint(self):
        tree = ast.parse((ROOT / "cm.py").read_text(encoding="utf-8-sig"))
        main_functions = [node for node in tree.body if isinstance(node, ast.FunctionDef) and node.name == "main"]
        guards = [
            node for node in tree.body
            if isinstance(node, ast.If)
            and isinstance(node.test, ast.Compare)
            and isinstance(node.test.left, ast.Name)
            and node.test.left.id == "__name__"
        ]
        self.assertEqual(len(main_functions), 1)
        self.assertEqual(len(guards), 1)

    def test_jsform_child_cleanup_ignores_deleted_panels(self):
        close_children = load_function_without_importing(
            Path(r"C:\Users\Pastor\Documents\JSForm\clsForm.py"),
            "_close_child_forms",
            {},
        )

        class LivePanel:
            def __init__(self):
                self.closed = False

            def IsBeingDeleted(self):
                return False

            def Close(self):
                self.closed = True

        class DeletedPanel:
            def IsBeingDeleted(self):
                raise RuntimeError("wrapped C/C++ object has been deleted")

        live = type("Child", (), {"FORM": LivePanel()})()
        deleted = type("Child", (), {"FORM": DeletedPanel()})()
        children = {"live": live, "deleted": deleted}

        close_children(object(), children)

        self.assertEqual(children, {})
        self.assertTrue(live.FORM.closed)

    def test_operational_modules_compile(self):
        with tempfile.TemporaryDirectory(prefix="churchmanager_compile_") as cache_dir:
            for relative in OPERATIONAL_MODULES:
                source = ROOT / relative
                with self.subTest(module=relative):
                    self.assertTrue(source.is_file(), f"Missing ChurchManager module: {relative}")
                    target = Path(cache_dir) / f"{source.stem}.pyc"
                    py_compile.compile(str(source), cfile=str(target), doraise=True)

    def test_report_week_calculation_boundaries(self):
        expected = {
            date(2026, 8, 1): 1,
            date(2026, 8, 2): 1,
            date(2026, 8, 8): 2,
            date(2026, 8, 9): 2,
            date(2026, 8, 30): 5,
            date(2026, 9, 1): 1,
            date(2026, 9, 7): 2,
        }
        function = load_function_without_importing(
            ROOT / "report_support.py", "get_week_of_month", {}
        )
        for value, week in expected.items():
            with self.subTest(value=value):
                self.assertEqual(function(value), week)

    def test_report_date_override(self):
        function = load_function_without_importing(
            ROOT / "report_support.py", "get_today",
            {"datetime": datetime, "date": date},
        )
        self.assertEqual(
            function({"testing": {"override_today": "2026-08-09"}}),
            date(2026, 8, 9),
        )

    def test_command_line_reports_have_guarded_entrypoints(self):
        for module_name in ("rptAnnouncement.py", "rptPrayers.py"):
            source = (ROOT / module_name).read_text(encoding="utf-8-sig")
            with self.subTest(module=module_name):
                self.assertIn('if __name__ == "__main__":', source)

    def test_plain_bulletin_output_uses_real_tabs(self):
        from bulletin_orders import render_plain_line

        self.assertEqual(
            render_plain_line("Opening Hymn", value="LSB 507", has_tab=True),
            "Opening Hymn\tLSB 507",
        )

    def test_bulletin_order_sql_uses_mysql_connector_markers(self):
        from bulletin_orders import _PortableCursor

        calls = []
        Cursor = type(
            "Cursor", (),
            {
                "__module__": "mysql.connector.cursor",
                "execute": lambda self, sql, values=(): calls.append((sql, values)),
            },
        )
        _PortableCursor(Cursor()).execute("SELECT ID FROM Sample WHERE ID=?", (7,))
        self.assertEqual(calls, [("SELECT ID FROM Sample WHERE ID=%s", (7,))])

    def test_sermon_docx_conversion_uses_expected_html_markers(self):
        function = load_function_without_importing(
            ROOT / "sermon2blogger.py", "convert_docx_to_text", {"Document": Document}
        )
        with tempfile.TemporaryDirectory(prefix="churchmanager_sermon_") as folder:
            source = Path(folder) / "sample.docx"
            output = Path(folder) / "sample.txt"
            document = Document()
            document.add_paragraph("Opening paragraph")
            quote_style = document.styles.add_style("Bible Quote", 1)
            quote_style.base_style = document.styles["Normal"]
            document.add_paragraph("For God so loved the world", style="Bible Quote")
            document.add_paragraph("")
            document.save(source)
            function(source, output)
            converted = output.read_text(encoding="utf-8")
            self.assertIn("Opening paragraph<br><br>", converted)
            self.assertIn("<blockquote>For God so loved the world</blockquote>", converted)


class TestChurchManagerConfiguration(unittest.TestCase):
    def test_main_menu_router_dispatches_forms_and_special_actions(self):
        from main_menu import MainMenuRouter

        opened = []
        special = []
        factory = type("Factory", (), {"open": lambda self, name: opened.append(name) or name})()
        router = MainMenuRouter(factory, {"lblBackupDB": lambda: special.append("backup")})
        self.assertEqual(router.dispatch("lblSermon"), "frmSermon")
        router.dispatch("lblBackupDB")
        self.assertEqual(opened, ["frmSermon"])
        self.assertEqual(special, ["backup"])

    def test_form_factory_preserves_parent_connection_and_controls(self):
        from form_factory import ChurchManagerFormFactory

        created = []

        class Form:
            def __init__(self, *arguments):
                created.append(arguments)
                self.shown = False

            def show(self):
                self.shown = True

        factory = ChurchManagerFormFactory(Form, "connection", "parent")
        form = factory.open("frmExample", ["Close"])
        self.assertEqual(created, [("parent", "connection", "frmExample", ["Close"])])
        self.assertTrue(form.shown)

    def test_form_factory_decorates_only_the_propers_editor(self):
        source = (ROOT / "form_factory.py").read_text(encoding="utf-8")
        self.assertIn('form_name == "frmPropers"', source)
        self.assertIn("LITURGICAL_COLOR_SWATCH", source)
        self.assertIn("liturgical_color_hex(color_field.GetValue())", source)

    def test_configuration_has_required_sections(self):
        config = load_json(ROOT / "churchmanager.json")
        self.assertIn("database_settings", config)
        self.assertIn("testing", config)
        self.assertTrue({"host", "database", "user", "credential_target"}.issubset(config["database_settings"]))
        self.assertNotIn("password", config["database_settings"])

    def test_testing_override_is_blank_or_iso_date(self):
        config = load_json(ROOT / "churchmanager.json")
        override = config.get("testing", {}).get("override_today")
        if override:
            datetime.strptime(override, "%Y-%m-%d")

    def test_test_mode_selects_an_isolated_database(self):
        from churchmanager_mode import resolve_database

        config = load_json(ROOT / "churchmanager.json")
        resolved = resolve_database(
            {"database": "ChurchDB", "test_mode": True, "password": None},
            config,
            lambda target: ("church", "not-a-real-password"),
        )
        self.assertEqual(resolved["database"], "ChurchDBTest")
        self.assertEqual(resolved["jsform_database"], "JSFormTest")
        self.assertEqual(resolved["server"], "127.0.0.1")
        self.assertEqual(resolved["credential_target"], "ChurchManager/LocalTestAdmin")
        self.assertNotEqual(
            resolved["database"].casefold(),
            config["database_settings"]["database"].casefold(),
        )

    def test_test_mode_refuses_production_database(self):
        from churchmanager_mode import resolve_database

        unsafe_config = {
            "database_settings": {"database": "ChurchDB"},
            "testing": {"host": "127.0.0.1", "database": "churchdb"},
        }
        with self.assertRaisesRegex(RuntimeError, "production database"):
            resolve_database(
                {"database": "ChurchDB", "test_mode": True}, unsafe_config,
                lambda target: ("church", "not-a-real-password"),
            )

    def test_test_mode_refuses_production_jsform_database(self):
        from churchmanager_mode import resolve_database

        unsafe_config = {
            "database_settings": {
                "database": "ChurchDB",
                "jsform_database": "JSForm",
            },
            "testing": {
                "host": "127.0.0.1",
                "database": "ChurchDBTest",
                "jsform_database": "jsform",
            },
        }
        with self.assertRaisesRegex(RuntimeError, "JSForm test database"):
            resolve_database(
                {"database": "ChurchDB", "test_mode": True}, unsafe_config,
                lambda target: ("church", "not-a-real-password"),
            )

    def test_database_password_comes_from_credential_store(self):
        from churchmanager_mode import resolve_database

        config = load_json(ROOT / "churchmanager.json")
        requested = []

        def fake_reader(target):
            requested.append(target)
            return "church", "stored-secret"

        resolved = resolve_database(
            {
                "database": "ChurchDB",
                "user": "church",
                "password": None,
                "test_mode": False,
                "jsform_database": None,
            },
            config,
            fake_reader,
        )
        self.assertEqual(requested, ["ChurchManager/Production"])
        self.assertEqual(resolved["password"], "stored-secret")

    def test_argument_parser_accepts_test_switch(self):
        import fnCMargParse

        parsed = fnCMargParse.CMargs(
            "ChurchManager", "test", ["database", "test_mode"], ["--test"]
        )
        self.assertTrue(parsed["test_mode"])

    def test_test_mode_title_is_applied_to_window_not_panel(self):
        source = (ROOT / "startup.py").read_text(encoding="utf-8-sig")
        self.assertIn("main_form.FRAME.SetTitle(", source)
        self.assertNotIn("main_form.FORM.SetTitle", source)

    def test_user_security_is_enabled_only_for_development_test_mode(self):
        from startup import security_enabled

        config = {
            "security": {"testing_enabled": True, "production_enabled": False}
        }
        self.assertTrue(security_enabled({"test_mode": True}, config))
        self.assertFalse(security_enabled({"test_mode": False}, config))

    def test_test_mode_uses_relaxed_password_length_only_at_composition_root(self):
        source = (ROOT / "startup.py").read_text(encoding="utf-8-sig")
        self.assertIn('minimum_length=4 if arguments["test_mode"] else 12', source)


class TestChurchManagerForms(unittest.TestCase):
    def test_propers_use_reusable_lectionary_system_and_optional_cycle(self):
        definition = load_json(FORMS / "frmPropers.json")["frmPropersFORM"]
        controls = definition["CONTROLS"]
        self.assertIn("LectionarySystemID", controls)
        self.assertEqual(controls["Cycle"]["choices"], ["A", "B", "C"])
        self.assertNotIn("Lectionary", controls)
        self.assertEqual(
            controls["LectionarySystemID"]["lookupchoices"]["name"],
            "tblLectionarySystem",
        )
        self.assertIn("frmLectionarySystem", definition["FORM"]["linkedform"])

    def test_participant_notifications_use_normalized_active_assignments(self):
        source = (ROOT / "fnSchedule.py").read_text(encoding="utf-8-sig")
        self.assertIn("tblServiceRole", source)
        self.assertIn("AssignmentStatus<>'DECLINED'", source)
        self.assertIn("p.Active=1", source)
        self.assertNotIn("tblSchedule", source)
        self.assertNotIn("p.Roles", source)
        self.assertNotIn("p.Schedule", source)

    def test_obsolete_jsform_database_structures_have_a_guarded_retirement(self):
        migration = (
            ROOT / "migrations" / "053_remove_obsolete_jsform_database_structures.sql"
        ).read_text(encoding="utf-8")
        runner = (ROOT / "run_churchdb_migrations.py").read_text(encoding="utf-8-sig")
        for table in ("tblOrderofService", "tblSchedule", "tblCheckList"):
            self.assertIn(f"DROP TABLE IF EXISTS {table}", migration)
        for column in ("OrderofService", "CheckListID", "Roles", "Schedule"):
            self.assertIn(f"DROP COLUMN IF EXISTS {column}", migration)
        self.assertIn("verify_obsolete_structure_conversion", runner)

    def test_worship_service_filters_propers_by_church_default_lectionary(self):
        source = (ROOT / "unified_worship_service_dialog.py").read_text(encoding="utf-8-sig")
        migration = (
            ROOT / "migrations" / "081_remove_lsb_lectionary_catalog.sql"
        ).read_text(encoding="utf-8")
        self.assertIn("PrimaryLectionaryEditionID", source)
        self.assertNotIn("PrimaryLectionarySystemID", source)
        self.assertIn("p.IsActive=1 AND e.IsActive=1", source)
        self.assertIn("def propers(self, church_id)", source)
        self.assertIn("DELETE FROM tblLectionarySystem", migration)
        self.assertIn("DELETE FROM tblServiceReadingSnapshot", migration)

    def test_user_administration_is_a_protected_main_menu_action(self):
        definition = next(iter(load_json(FORMS / "frmMain.json").values()))
        control = definition["CONTROLS"]["lblUsers"]
        self.assertEqual(control["security"]["invoke"], "security.users.manage")
        from main_menu import MENU_CONTROLS
        self.assertIn("lblUsers", MENU_CONTROLS)

    def test_report_designer_is_a_separately_protected_main_menu_action(self):
        definition = next(iter(load_json(FORMS / "frmMain.json").values()))
        control = definition["CONTROLS"]["lblReportDesigner"]
        self.assertEqual(control["security"]["invoke"], "reports.design")
        from main_menu import MENU_CONTROLS
        self.assertIn("lblReportDesigner", MENU_CONTROLS)
        source = (ROOT / "cm.py").read_text(encoding="utf-8-sig")
        self.assertIn('case "lblReportDesigner":', source)
        self.assertIn(
            "open_directory_designer(authorization=context.authorization)", source
        )

    def test_every_main_menu_action_declares_its_catalog_permission(self):
        from main_menu import MENU_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS

        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(set(MENU_CONTROLS), set(MAIN_MENU_PERMISSIONS))
        for control_name, permission in MAIN_MENU_PERMISSIONS.items():
            self.assertEqual(
                controls[control_name]["security"]["invoke"], permission,
                control_name,
            )

    def test_main_menu_dispatch_rechecks_permissions(self):
        source = (ROOT / "cm.py").read_text(encoding="utf-8-sig")
        self.assertIn("MAIN_MENU_PERMISSIONS[select]", source)
        self.assertIn("context.authorization.require(", source)

    def test_main_menu_provides_current_user_password_change_and_logout(self):
        from main_menu import SESSION_CONTROLS

        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(SESSION_CONTROLS, {"lblHelp", "lblChangePassword", "lblLogout"})
        self.assertIn("lblCurrentUser", controls)
        self.assertIn("LOGOUT", (ROOT / "cm.py").read_text(encoding="utf-8-sig"))

    def test_chart_of_accounts_is_a_protected_jsform_route(self):
        from main_menu import FORM_ROUTES

        self.assertEqual(FORM_ROUTES["lblAccountingAccounts"], "frmAccountingAccount")
        definition = load_json(FORMS / "frmAccountingAccount.json")[
            "frmAccountingAccountFORM"
        ]
        security = definition["FORM"]["security"]
        self.assertEqual(
            set(security.values()), {"accounting.master_data.manage"}
        )
        fields = definition["FORM"]["table"]["fields"]
        self.assertEqual(fields, ["*"])

    def test_funds_are_a_protected_jsform_route(self):
        from main_menu import FORM_ROUTES

        self.assertEqual(FORM_ROUTES["lblAccountingFunds"], "frmAccountingFund")
        definition = load_json(FORMS / "frmAccountingFund.json")[
            "frmAccountingFundFORM"
        ]
        self.assertEqual(
            set(definition["FORM"]["security"].values()),
            {"accounting.master_data.manage"},
        )
        controls = definition["CONTROLS"]
        self.assertIn("WITH_DONOR_RESTRICTIONS", controls["NetAssetClass"]["choices"])
        self.assertIn("BoardDesignated", controls)

    def test_functions_are_a_protected_jsform_route(self):
        from main_menu import FORM_ROUTES

        self.assertEqual(
            FORM_ROUTES["lblAccountingFunctions"], "frmAccountingFunction"
        )
        definition = load_json(FORMS / "frmAccountingFunction.json")[
            "frmAccountingFunctionFORM"
        ]
        self.assertEqual(
            set(definition["FORM"]["security"].values()),
            {"accounting.master_data.manage"},
        )
        choices = definition["CONTROLS"]["FunctionClass"]["choices"]
        self.assertEqual(
            choices, ["PROGRAM", "MANAGEMENT_GENERAL", "FUNDRAISING"]
        )

    def test_fiscal_years_and_periods_protect_status_overrides(self):
        from main_menu import FORM_ROUTES

        expected = {
            "lblAccountingYears": "frmAccountingFiscalYear",
            "lblAccountingPeriods": "frmAccountingFiscalPeriod",
        }
        for control_name, form_name in expected.items():
            self.assertEqual(FORM_ROUTES[control_name], form_name)
            definition = load_json(FORMS / (form_name + ".json"))[
                form_name + "FORM"
            ]
            self.assertEqual(
                definition["CONTROLS"]["Status"]["security"]["edit"],
                "accounting.periods.override",
            )
            self.assertEqual(
                set(definition["FORM"]["security"].values()),
                {"accounting.master_data.manage"},
            )
        main_controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        box = main_controls["FundAccountingBox"]
        left, top = box["posch"]
        width, height = box["sizech"]
        for name in ("lblAccountingYears", "lblAccountingPeriods"):
            x, y = main_controls[name]["posch"]
            self.assertTrue(left < x < left + width, name)
            self.assertTrue(top < y < top + height, name)

    def test_transaction_entry_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS

        self.assertIn("lblAccountingTransactions", SPECIAL_CONTROLS)
        self.assertEqual(
            MAIN_MENU_PERMISSIONS["lblAccountingTransactions"],
            "accounting.transactions.create",
        )
        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(
            controls["lblAccountingTransactions"]["security"]["invoke"],
            "accounting.transactions.create",
        )
        source = (ROOT / "cm.py").read_text(encoding="utf-8-sig")
        self.assertIn("show_accounting_draft_entry(", source)

    def test_transaction_review_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS

        self.assertIn("lblAccountingReview", SPECIAL_CONTROLS)
        self.assertEqual(
            MAIN_MENU_PERMISSIONS["lblAccountingReview"],
            "accounting.transactions.approve",
        )
        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(
            controls["lblAccountingReview"]["security"]["invoke"],
            "accounting.transactions.approve",
        )

    def test_transaction_posting_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS

        self.assertIn("lblAccountingPosting", SPECIAL_CONTROLS)
        self.assertEqual(
            MAIN_MENU_PERMISSIONS["lblAccountingPosting"],
            "accounting.transactions.post",
        )
        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(
            controls["lblAccountingPosting"]["security"]["invoke"],
            "accounting.transactions.post",
        )

    def test_posted_register_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS
        self.assertIn("lblAccountingRegister", SPECIAL_CONTROLS)
        self.assertEqual(MAIN_MENU_PERMISSIONS["lblAccountingRegister"],
                         "accounting.transactions.view")
        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(controls["lblAccountingRegister"]["security"]["invoke"],
                         "accounting.transactions.view")

    def test_trial_balance_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS
        self.assertIn("lblAccountingTrialBalance", SPECIAL_CONTROLS)
        self.assertEqual(MAIN_MENU_PERMISSIONS["lblAccountingTrialBalance"], "accounting.reports.run")
        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(controls["lblAccountingTrialBalance"]["security"]["invoke"], "accounting.reports.run")

    def test_financial_position_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS
        self.assertIn("lblAccountingPosition",SPECIAL_CONTROLS)
        self.assertEqual(MAIN_MENU_PERMISSIONS["lblAccountingPosition"],"accounting.reports.run")
        controls=load_json(FORMS/"frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(controls["lblAccountingPosition"]["security"]["invoke"],"accounting.reports.run")

    def test_statement_of_activities_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS
        self.assertIn("lblAccountingActivities",SPECIAL_CONTROLS)
        self.assertEqual(MAIN_MENU_PERMISSIONS["lblAccountingActivities"],"accounting.reports.run")
        controls=load_json(FORMS/"frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(controls["lblAccountingActivities"]["security"]["invoke"],"accounting.reports.run")

    def test_bank_accounts_are_protected_jsform_master_data(self):
        from main_menu import FORM_ROUTES
        from permission_catalog import MAIN_MENU_PERMISSIONS
        self.assertEqual(FORM_ROUTES["lblAccountingBankAccounts"],"frmAccountingBankAccount")
        self.assertEqual(MAIN_MENU_PERMISSIONS["lblAccountingBankAccounts"],"accounting.master_data.manage")
        form=load_json(FORMS/"frmAccountingBankAccount.json")["frmAccountingBankAccountFORM"]
        self.assertEqual(form["FORM"]["table"]["name"],"tblAccountingBankAccount")
        lookup=form["CONTROLS"]["AccountID"]["lookupchoices"]
        self.assertIn("AccountType = 'ASSET'",lookup["condition"])
        self.assertIn("last four digits",form["CONTROLS"]["AccountLastFour"]["tooltip"])
        controls=load_json(FORMS/"frmMain.json")["frmMainFORM"]["CONTROLS"]
        box=controls["DesignersBox"]; top=box["posch"][1]; bottom=top+box["sizech"][1]
        self.assertTrue(top < controls["lblAccountingActivities"]["posch"][1] < bottom)

    def test_payees_are_protected_jsform_master_data(self):
        from main_menu import FORM_ROUTES
        from permission_catalog import MAIN_MENU_PERMISSIONS

        self.assertEqual(FORM_ROUTES["lblAccountingPayees"], "frmAccountingPayee")
        self.assertEqual(
            MAIN_MENU_PERMISSIONS["lblAccountingPayees"],
            "accounting.master_data.manage",
        )
        form = load_json(FORMS / "frmAccountingPayee.json")["frmAccountingPayeeFORM"]
        self.assertEqual(form["FORM"]["table"]["name"], "tblAccountingPayee")
        self.assertEqual(form["CONTROLS"]["OrganizationID"]["type"], "ComboBox")
        self.assertEqual(form["CONTROLS"]["ContactData"]["type"], "MultiLine")

    def test_bank_import_is_a_protected_special_workflow(self):
        from main_menu import SPECIAL_CONTROLS
        from permission_catalog import MAIN_MENU_PERMISSIONS

        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertIn("lblAccountingBankImport", SPECIAL_CONTROLS)
        self.assertEqual(
            MAIN_MENU_PERMISSIONS["lblAccountingBankImport"],
            "accounting.reconciliation.manage",
        )
        self.assertEqual(
            controls["lblAccountingBankImport"]["security"]["invoke"],
            "accounting.reconciliation.manage",
        )

    def test_forms_match_jsform_canonical_schema(self):
        from jsonschema import validate

        schema_path = Path(r"C:\Users\Pastor\Documents\JSForm\schema\unified_schema.json")
        schema = load_json(schema_path)
        for path in sorted(FORMS.glob("*.json")):
            with self.subTest(form=path.name):
                validate(instance=load_json(path), schema=schema)

    def test_all_application_forms_are_valid_json(self):
        files = sorted(FORMS.glob("*.json"))
        self.assertGreater(files, [], "No ChurchManager form definitions found")
        for path in files:
            with self.subTest(form=path.name):
                load_json(path)

    def test_form_filename_matches_top_level_definition(self):
        for path in sorted(FORMS.glob("*.json")):
            data = load_json(path)
            expected = f"{path.stem}FORM".casefold()
            with self.subTest(form=path.name):
                self.assertIn(expected, {key.casefold() for key in data})

    def test_forms_have_form_and_controls_sections(self):
        for path in sorted(FORMS.glob("*.json")):
            data = load_json(path)
            with self.subTest(form=path.name):
                self.assertEqual(len(data), 1, "Expected one top-level form definition")
                definition = next(iter(data.values()))
                self.assertIsInstance(definition.get("FORM"), dict)
                self.assertIsInstance(definition.get("CONTROLS"), dict)

    def test_open_file_actions_reference_local_controls(self):
        for path in sorted(FORMS.glob("*.json")):
            definition = next(iter(load_json(path).values()))
            controls = definition.get("CONTROLS", {})
            for control_key, control in controls.items():
                action = control.get("action") if isinstance(control, dict) else None
                if isinstance(action, list) and len(action) > 1 and action[0] == "openfile":
                    with self.subTest(form=path.name, control=control_key):
                        self.assertIn(action[1], controls)

    def test_linked_form_actions_reference_churchmanager_form_files(self):
        available = {path.stem.casefold() for path in FORMS.glob("*.json")}
        for path in sorted(FORMS.glob("*.json")):
            definition = next(iter(load_json(path).values()))
            for control_key, control in definition.get("CONTROLS", {}).items():
                action = control.get("action") if isinstance(control, dict) else None
                if isinstance(action, list) and len(action) > 1 and action[0] == "openlinkedform":
                    with self.subTest(form=path.name, control=control_key):
                        self.assertIn(action[1].casefold(), available)

    def test_file_picker_directories_have_section_and_key(self):
        for path in sorted(FORMS.glob("*.json")):
            definition = next(iter(load_json(path).values()))
            for control_key, control in definition.get("CONTROLS", {}).items():
                if isinstance(control, dict) and control.get("type") == "FilePickerCtrl":
                    directory = control.get("directory")
                    with self.subTest(form=path.name, control=control_key):
                        self.assertIsInstance(directory, list)
                        self.assertEqual(len(directory), 2)
                        self.assertTrue(all(isinstance(value, str) and value for value in directory))

    def test_bound_main_menu_controls_exist(self):
        from main_menu import MENU_CONTROLS

        bound = set(MENU_CONTROLS)
        main = next(iter(load_json(FORMS / "frmMain.json").values()))
        controls = set(main["CONTROLS"])
        self.assertGreater(bound, set(), "No ChurchManager main-menu bindings found")
        self.assertEqual(bound - controls, set(), "cm.py binds controls missing from frmMain.json")

    def test_main_menu_has_compact_four_column_dashboard(self):
        main = next(iter(load_json(FORMS / "frmMain.json").values()))
        controls = main["CONTROLS"]
        expected = {
            "ChurchBox": "Service Planning",
            "MemberBox": "People and Congregation",
            "ServiceBox": "Worship Resources",
            "ReportBox": "Reports and Design",
            "UtilitiesBox": "ChurchManager Settings",
            "SessionBox": "Current User",
            "BulletinBox": "Accounting - Daily Work",
            "DesignersBox": "Accounting - Reports",
            "FundAccountingBox": "Accounting - Setup and Close",
        }
        self.assertEqual(
            {name: controls[name]["label"] for name in expected}, expected,
        )
        self.assertEqual({controls[name]["posch"][0] for name in expected}, {1, 14, 27, 40})

    def test_main_menu_groups_work_by_usage(self):
        controls = load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"]
        self.assertEqual(controls["ServiceBox"]["label"], "Worship Resources")
        self.assertEqual(controls["ChurchBox"]["label"], "Service Planning")
        planning_items = [
            "lblService", "lblWeeklyBulletinOrder", "lblServiceSchedule",
            "lblNotifyParticipants", "lblSundayPrayers", "lblAnnouncements",
            "lblGenerateOS",
        ]
        self.assertEqual(
            sorted(controls[name]["posch"][1] for name in planning_items),
            list(range(2, 9)),
        )
        resource_items = [
            "lblOS", "lblCheckList", "lblPropers", "lblSermon", "lblHymnal",
            "lblHymn", "lblParticipant", "lblSchedule", "lblPrayers",
            "lblWorshipPositions", "lblAnnouncement", "lblAttendanceEvent",
        ]
        self.assertEqual(
            sorted(controls[name]["posch"][1] for name in resource_items),
            list(range(12, 24)),
        )
        self.assertTrue(all(controls[name]["label"] == controls[name]["label"].strip()
                            for name in resource_items))
        from main_menu import FORM_ROUTES
        self.assertEqual(FORM_ROUTES["lblHymnal"], "frmHymnal")
        self.assertEqual(FORM_ROUTES["lblHymn"], "frmHymn")

    def test_removed_financial_features_are_not_exposed(self):
        removed_forms = {
            "frmBudget", "frmFund", "frmLedger", "frmCheckRegister",
            "frmGivingRegister", "frmEnvelope", "frmDonor", "frmDonorGift",
            "frmPostCheck", "frmPostGiving",
        }
        removed_controls = {
            "AccountingBox", "DonorBox", "lblChartofAccounts", "lblBudget",
            "lblLedger", "lblCheckRegister", "lblGivingRegister", "lblEnvelope",
            "lblDonor", "lblDonorGift",
        }
        removed_reports = {"CFCA01", "CFCR01", "CFGR01", "CMDN01", "CMDN02"}

        from main_menu import FORM_ROUTES

        main = next(iter(load_json(FORMS / "frmMain.json").values()))
        report_filter = (ROOT / "report_access.py").read_text(encoding="utf-8")

        self.assertTrue(removed_forms.isdisjoint({path.stem for path in FORMS.glob("*.json")}))
        self.assertTrue(removed_controls.isdisjoint(main["CONTROLS"]))
        self.assertTrue(removed_forms.isdisjoint(FORM_ROUTES.values()))
        self.assertTrue(removed_reports.isdisjoint({path.stem for path in VISUAL_REPORTS.glob("*.json")}))
        for report_code in removed_reports:
            self.assertIn(report_code, report_filter)


class TestChurchManagerReportAssets(unittest.TestCase):
    def test_report_menu_keeps_church_readable_and_focuses_report_grid(self):
        source = (ROOT / "cm.py").read_text(encoding="utf-8")
        report_case = source.split('case "lblReports":', 1)[1].split(
            'case "lblReportDesigner":', 1,
        )[0]
        self.assertIn('enable_button("ChurchID")', report_case)
        self.assertIn('CONTROLID["ReportID"].SetFocus()', report_case)

    def test_churchmanager_has_no_lime_report_runtime_assets(self):
        patterns = ROOT / "LimeReportPattern"
        self.assertFalse(any(patterns.glob("*.lrxml")))
        self.assertFalse(any(patterns.glob("*.lrsml")))
        source = (ROOT / "report_service.py").read_text(encoding="utf-8")
        self.assertNotIn("LimeReportProcess", source)
        self.assertNotIn("prepare_lime_report_template", source)

    def test_church_form_exposes_database_logo_image_picker(self):
        church = next(iter(load_json(FORMS / "frmChurch.json").values()))
        logo = church["CONTROLS"]["Logo"]
        self.assertEqual(logo["type"], "ImagePickerCtrl")
        self.assertEqual(logo["name"], "Logo")
        self.assertGreater(logo["maxbytes"], 0)


class TestNonAccountingTestDataset(unittest.TestCase):
    def test_safe_report_views_filter_unlisted_and_non_directory_records(self):
        source = (ROOT / "migrations" / "017_add_nonaccounting_report_views.sql").read_text(
            encoding="utf-8-sig"
        )
        for view in (
            "rpt_person_contact", "rpt_person_address",
            "rpt_family_contact", "rpt_family_address",
        ):
            start = source.index(f"VIEW {view}")
            statement = source[start:source.index(";", start)]
            with self.subTest(view=view):
                self.assertIn("WHERE Unlisted=0", statement)
        start = source.index("VIEW rpt_directory_family")
        statement = source[start:source.index(";", start)]
        self.assertIn("WHERE Directory=1", statement)

    def test_seeder_is_local_test_only_and_excludes_accounting_writes(self):
        source = (ROOT / "seed_nonaccounting_test_data.py").read_text(
            encoding="utf-8-sig"
        )
        self.assertIn('database.casefold() != "churchdbtest"', source)
        self.assertIn("LOCAL_HOSTS", source)
        self.assertIn('parser.add_argument("--apply"', source)
        self.assertIn("connection.rollback()", source)
        self.assertNotIn("INSERT INTO tblAccounting", source)
        self.assertNotIn("Synthetic Test Hymnal", source)

    def test_seeder_uses_fictional_church_logo_and_runtime_password_prompt(self):
        source = (ROOT / "seed_nonaccounting_test_data.py").read_text(
            encoding="utf-8-sig"
        )
        self.assertIn('CHURCH_NAME = "Reformation Lutheran Church"', source)
        self.assertIn("Reformation-Lutheran-Church-Test-Logo.png", source)
        self.assertIn("getpass.getpass", source)
        self.assertNotIn("Password123", source)

    def test_report_permission_migration_categorizes_every_current_report(self):
        source = (ROOT / "migrations" / "015_add_report_permissions.sql").read_text(
            encoding="utf-8-sig"
        )
        self.assertIn("RequiredPermissionID", source)
        self.assertIn("fk_reports_required_permission", source)
        self.assertIn("ALTER TABLE tblReports MODIFY RequiredPermissionID", source)
        for code in (
            "CMAS01", "CMAT01", "CMBATCH00", "CMDO01",
            "CMHU01", "CMHU02", "CMHU03", "CMHU04", "CMJR01",
            "CMMD01", "CMMI01", "CMMI02", "CMMI03", "CMML01",
            "CMML02", "CMPA01", "CMPE01", "CMPH02", "CMPJ01",
            "CMPJ02", "CMPJ03", "CMPJ04", "CMPR01", "CMRP01",
            "CMSM01", "CMWP01", "CMWS01",
        ):
            with self.subTest(report=code):
                self.assertIn("'{}'".format(code), source)

    def test_report_category_roles_can_open_the_report_picker(self):
        source = (
            ROOT / "migrations" / "016_grant_report_screen_access.sql"
        ).read_text(encoding="utf-8-sig")
        self.assertIn("p.Name='reports.run'", source)
        for role in ("Master Administrator", "Pastor/Staff", "Volunteer", "Auditor"):
            self.assertIn("'{}'".format(role), source)

    def test_obsolete_enhancement_tracker_is_retired(self):
        migration = (
            ROOT / "migrations" / "066_retire_enhancement_tracker.sql"
        ).read_text(encoding="utf-8-sig")
        self.assertIn("DROP VIEW IF EXISTS rpt_enhancement", migration)
        self.assertIn("WHERE Report='CMEN01'", migration)
        self.assertIn("DROP TABLE IF EXISTS tblEnhancement", migration)
        self.assertNotIn("lblEnhancements", load_json(FORMS / "frmMain.json")["frmMainFORM"]["CONTROLS"])
        self.assertFalse((ROOT / "visual_reports" / "definitions" / "CMEN01.json").exists())

    def test_report_screen_uses_compact_explicit_layout(self):
        report_form = next(iter(load_json(FORMS / "frmReports.json").values()))
        form = report_form["FORM"]
        controls = report_form["CONTROLS"]
        self.assertEqual(form["layout"]["type"], "responsive")
        self.assertEqual(controls["ChurchID"]["value"], "Reformation Lutheran Church")
        self.assertLessEqual(controls["ChurchID"]["sizech"][1], 2)
        self.assertGreaterEqual(controls["ParameterBox"]["sizech"][1], 24)
        self.assertEqual(controls["ParameterBox"]["layout"]["column_span"], 2)
        self.assertGreater(
            controls["btnRun"]["layout"]["row"],
            controls["ProjectID"]["posch"][1],
        )


if __name__ == "__main__":
    unittest.main()
