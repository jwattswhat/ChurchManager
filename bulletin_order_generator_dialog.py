"""Preview and export a service's structured bulletin order."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches
import wx

from bulletin_orders import BulletinOrderGenerator


OUTPUTS = ("Plain Text", "HTML", "Word Document")


def save_word_order(path, rendered):
    document = Document()
    for item in rendered["lines"]:
        paragraph = document.add_paragraph()
        if item["tab_position"] is not None:
            alignments = {
                "LEFT": WD_TAB_ALIGNMENT.LEFT, "CENTER": WD_TAB_ALIGNMENT.CENTER,
                "RIGHT": WD_TAB_ALIGNMENT.RIGHT, "DECIMAL": WD_TAB_ALIGNMENT.DECIMAL,
            }
            leaders = {"NONE": WD_TAB_LEADER.SPACES, "DOTS": WD_TAB_LEADER.DOTS,
                       "LINE": WD_TAB_LEADER.LINES}
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(float(item["tab_position"])),
                alignments.get(item["tab_alignment"], WD_TAB_ALIGNMENT.RIGHT),
                leaders.get(item["tab_leader"], WD_TAB_LEADER.SPACES),
            )
        if item["indent"]:
            paragraph.paragraph_format.left_indent = Inches(0.25 * int(item["indent"]))
        label = paragraph.add_run(item["label"] or "")
        label.bold = item["label_bold"]
        label.italic = item["italic"]
        right = item["value"] or item["reference"]
        if right:
            paragraph.add_run("\t" if item["tab_position"] is not None else " ")
            value = paragraph.add_run(str(right))
            value.bold = item["value_bold"]
            value.italic = item["italic"]
    document.save(path)


class PrepareBulletinOrderDialog(wx.Dialog):
    def __init__(self, parent, connection):
        super().__init__(parent, title="Prepare Bulletin Order", size=(1020, 700),
                         style=wx.DEFAULT_DIALOG_STYLE | wx.RESIZE_BORDER)
        self.generator = BulletinOrderGenerator(connection)
        self.service_rows = self.generator.services()
        self.template_rows = [row for row in self.generator.repository.templates() if row[3]]
        self.rendered = None
        self._build()
        self._load_choices()
        self.CentreOnParent()

    def _build(self):
        panel = wx.Panel(self)
        outer = wx.BoxSizer(wx.VERTICAL)
        selectors = wx.FlexGridSizer(cols=2, vgap=8, hgap=12)
        selectors.AddGrowableCol(1, 1)
        self.service = wx.Choice(panel)
        self.template = wx.Choice(panel)
        self.output = wx.Choice(panel, choices=OUTPUTS)
        self.output.SetSelection(0)
        for label, control in (("Service", self.service), ("Bulletin-order template", self.template),
                               ("Output format", self.output)):
            selectors.Add(wx.StaticText(panel, label=label), 0, wx.ALIGN_CENTER_VERTICAL)
            selectors.Add(control, 1, wx.EXPAND)
        outer.Add(selectors, 0, wx.EXPAND | wx.ALL, 12)

        self.grid = wx.ListCtrl(panel, style=wx.LC_REPORT | wx.LC_SINGLE_SEL)
        for label, width in (("Order", 65), ("Bulletin output", 510), ("Type", 110), ("Status", 150)):
            self.grid.AppendColumn(label, width=width)
        outer.Add(self.grid, 1, wx.EXPAND | wx.LEFT | wx.RIGHT, 12)
        self.status = wx.StaticText(panel, label="Choose a service and template, then select Preview.")
        outer.Add(self.status, 0, wx.EXPAND | wx.ALL, 12)

        buttons = wx.BoxSizer(wx.HORIZONTAL)
        for label, handler in (("Preview", self.on_preview), ("Copy Output", self.on_copy),
                               ("Save Output", self.on_save)):
            button = wx.Button(panel, label=label)
            button.Bind(wx.EVT_BUTTON, handler)
            buttons.Add(button, 0, wx.RIGHT, 8)
        buttons.AddStretchSpacer()
        close = wx.Button(panel, wx.ID_CLOSE, "Close")
        close.Bind(wx.EVT_BUTTON, lambda _event: self.EndModal(wx.ID_CLOSE))
        buttons.Add(close)
        outer.Add(buttons, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.BOTTOM, 12)
        panel.SetSizer(outer)
        self.service.Bind(wx.EVT_CHOICE, self.on_service)

    def _load_choices(self):
        self.service.Clear()
        self.template.Clear()
        for row in self.service_rows:
            when = row[1].strftime("%m/%d/%Y %I:%M %p") if hasattr(row[1], "strftime") else str(row[1])
            title = row[2] or "Service"
            self.service.Append(f"{when} — {title}")
        for row in self.template_rows:
            self.template.Append(row[1] + (" (Starter)" if row[4] else " (Customized)"))
        if self.service_rows:
            self.service.SetSelection(0)
            self.on_service()

    def on_service(self, _event=None):
        index = self.service.GetSelection()
        if index < 0:
            return
        suggested = self.generator.suggested_template_id(self.service_rows[index][0])
        template_index = next((i for i, row in enumerate(self.template_rows) if row[0] == suggested), 0)
        if self.template_rows:
            self.template.SetSelection(template_index)

    def on_preview(self, _event=None):
        service_index, template_index = self.service.GetSelection(), self.template.GetSelection()
        if service_index < 0 or template_index < 0:
            wx.MessageBox("Choose both a service and a bulletin-order template.", "Selection needed",
                          wx.OK | wx.ICON_INFORMATION, self)
            return
        try:
            self.rendered = self.generator.render(
                self.template_rows[template_index][0], self.service_rows[service_index][0]
            )
            self.grid.DeleteAllItems()
            missing = 0
            for index, item in enumerate(self.rendered["lines"]):
                row = self.grid.InsertItem(index, str(item["sequence"]))
                right = item["value"] or item["reference"] or ""
                display = item["label"] + (("    " + str(right)) if right else "")
                self.grid.SetItem(row, 1, display)
                self.grid.SetItem(row, 2, item["type"].title())
                status = f"Missing {item['value_key']}" if item["missing"] else "Ready"
                self.grid.SetItem(row, 3, status)
                if item["missing"]:
                    self.grid.SetItemTextColour(row, wx.RED)
                    missing += 1
            self.status.SetLabel(
                f"{len(self.rendered['lines'])} line(s); "
                + (f"{missing} missing value(s) must be resolved." if missing else "ready to copy or save.")
            )
        except Exception as error:
            wx.MessageBox(str(error), "Unable to prepare bulletin order", wx.OK | wx.ICON_ERROR, self)

    def _ensure_preview(self):
        if self.rendered is None:
            self.on_preview()
        return self.rendered is not None

    def on_copy(self, _event):
        if not self._ensure_preview():
            return
        selected = self.output.GetStringSelection()
        if selected == "Word Document":
            wx.MessageBox("Use Save Output to create a Word document.", "Word output",
                          wx.OK | wx.ICON_INFORMATION, self)
            return
        text = self.rendered["plain_text"] if selected == "Plain Text" else self.rendered["html"]
        if wx.TheClipboard.Open():
            try:
                wx.TheClipboard.SetData(wx.TextDataObject(text))
            finally:
                wx.TheClipboard.Close()
            self.status.SetLabel(f"{selected} copied to the clipboard.")

    def on_save(self, _event):
        if not self._ensure_preview():
            return
        selected = self.output.GetStringSelection()
        extensions = {"Plain Text": ("txt", "Text files (*.txt)|*.txt"),
                      "HTML": ("html", "HTML files (*.html)|*.html"),
                      "Word Document": ("docx", "Word documents (*.docx)|*.docx")}
        extension, wildcard = extensions[selected]
        reports = Path(__file__).resolve().parent / "Reports"
        reports.mkdir(exist_ok=True)
        dialog = wx.FileDialog(self, "Save Bulletin Order", defaultDir=str(reports),
                               defaultFile=f"Bulletin Order.{extension}", wildcard=wildcard,
                               style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT)
        try:
            if dialog.ShowModal() != wx.ID_OK:
                return
            path = Path(dialog.GetPath())
            if selected == "Word Document":
                save_word_order(path, self.rendered)
            else:
                value = self.rendered["plain_text"] if selected == "Plain Text" else self.rendered["html"]
                path.write_text(value, encoding="utf-8")
            self.status.SetLabel(f"Saved {path.name}")
        finally:
            dialog.Destroy()


def show_prepare_bulletin_order(parent, connection):
    dialog = PrepareBulletinOrderDialog(parent, connection)
    try:
        return dialog.ShowModal()
    finally:
        dialog.Destroy()
