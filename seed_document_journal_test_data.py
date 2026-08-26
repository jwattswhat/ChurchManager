"""Seed a sample document and fictional congregational journal entries."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

import mariadb
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from credential_store import read_credential


ROOT = Path(__file__).resolve().parent
CHURCH_NAME = "Reformation Lutheran Church"
MARKER = "CMTEST: document-journal sample"
LOCAL_HOSTS = {"localhost", "127.0.0.1", "::1"}


def create_sample_docx(path):
    """Create a small, valid Word document for open-file testing."""
    document = Document()
    section = document.sections[0]
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Reformation Lutheran Church")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(46, 116, 181)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Sample Congregational Planning Document")
    run.bold = True
    run.font.size = Pt(13)

    document.add_paragraph(
        "This fictional Word document verifies that ChurchManager opens DOCX "
        "files with the application registered in Windows."
    )
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run("Test notes")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(46, 116, 181)
    document.add_paragraph("No real congregation or member information is included.")
    document.add_paragraph("Created for the ChurchDBTest document catalog.")
    document.core_properties.title = "Sample Congregational Planning Document"
    document.core_properties.subject = "ChurchManager fictional test data"
    document.core_properties.author = "ChurchManager Test System"
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def settings():
    """Return guarded ChurchDBTest connection settings and credentials."""
    config = json.loads((ROOT / "churchmanager.json").read_text(encoding="utf-8-sig"))
    testing = config["testing"]
    host = str(testing["host"])
    database = str(testing["database"])
    if host not in LOCAL_HOSTS or database.casefold() != "churchdbtest":
        raise RuntimeError("Safety stop: seeding is restricted to local ChurchDBTest.")
    username, password = read_credential(testing["credential_target"])
    return testing, username, password


def scalar(cursor, sql, values=()):
    """Return the first column from a one-row query."""
    cursor.execute(sql, values)
    row = cursor.fetchone()
    return row[0] if row else None


def main():
    """Preview or apply the idempotent fictional document and journal fixture."""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--apply", action="store_true", help="commit the sample records")
    args = parser.parse_args()
    testing, username, password = settings()
    connection = mariadb.connect(
        host=testing["host"], port=int(testing.get("port", 3306)),
        database=testing["database"], user=username, password=password,
        autocommit=False,
    )
    password = ""
    cursor = connection.cursor()
    try:
        church_id = scalar(cursor, "SELECT ID FROM tblChurch WHERE Church=?", (CHURCH_NAME,))
        if not church_id:
            raise RuntimeError("The Reformation Lutheran Church test record is missing.")
        sample_path = ROOT / "Documents" / "Sample Congregational Document.txt"
        sample_docx_path = ROOT / "Documents" / "Test.Document.70.docx"
        if not sample_path.is_file():
            raise RuntimeError("The sample congregational document is missing.")

        print("target", testing["database"])
        print("church", CHURCH_NAME, church_id)
        print("document", sample_path)
        print("word_document", sample_docx_path)
        print("journal_entries", 3)
        if not args.apply:
            connection.rollback()
            print("No changes made. Re-run with --apply after reviewing this preview.")
            return 2

        create_sample_docx(sample_docx_path)

        cursor.execute("DELETE FROM tblDocument WHERE Note=?", (MARKER,))
        cursor.execute(
            "INSERT INTO tblDocument "
            "(ChurchID,Title,Document,Date,DocumentType,Description,Note) "
            "VALUES (?,?,?,?,?,?,?)",
            (
                church_id, "Sample Congregational Planning Note",
                str(sample_path.relative_to(ROOT)), "2026-08-26", "Planning",
                "Fictional sample file for testing the congregational document catalog.",
                MARKER,
            ),
        )
        cursor.execute("DELETE FROM tblJournal WHERE Note LIKE ?", (MARKER + "%",))
        entries = (
            ("Council approved the autumn ministry calendar", "2026-08-04 19:00:00", "2026-08-04 20:30:00"),
            ("Property committee reviewed seasonal maintenance", "2026-08-11 18:30:00", "2026-08-11 19:15:00"),
            ("Congregational fellowship planning follow-up", "2026-08-18 10:00:00", "2026-08-18 10:45:00"),
        )
        for number, (event, start, end) in enumerate(entries, 1):
            cursor.execute(
                "INSERT INTO tblJournal (ChurchID,Event,Complete,StartDate,EndDate,Note) "
                "VALUES (?,?,1,?,?,?)",
                (church_id, event, start, end, f"{MARKER} {number}"),
            )
        connection.commit()
        print("applied", True)
        print("documents", 1)
        print("word_document_created", sample_docx_path.is_file())
        print("journal_entries", len(entries))
        return 0
    except Exception:
        connection.rollback()
        raise
    finally:
        cursor.close()
        connection.close()


if __name__ == "__main__":
    raise SystemExit(main())
